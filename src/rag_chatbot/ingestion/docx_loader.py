from pathlib import Path

from docx import Document as WordDocument
from langchain_core.documents import Document

from rag_chatbot.ingestion.base_loader import BaseDocumentLoader


class DOCXDocumentLoader(BaseDocumentLoader):
    """Load text from Word documents."""

    def load(self, file_path: Path) -> list[Document]:
        word_document = WordDocument(file_path)

        paragraphs = [
            paragraph.text.strip()
            for paragraph in word_document.paragraphs
            if paragraph.text.strip()
        ]

        if not paragraphs:
            return []

        return [
            Document(
                page_content="\n\n".join(paragraphs),
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": file_path.suffix.lower(),
                },
            )
        ]