import json
from pathlib import Path

from docx import Document as WordDocument
from openpyxl import Workbook
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


PROJECT_ROOT = Path(__file__).resolve().parents[1]

RAW_DATA_DIR = PROJECT_ROOT / "data" / "raw"
EVALUATION_DIR = PROJECT_ROOT / "data" / "evaluation"


def ensure_directories() -> None:
    """Create dataset directories if they do not already exist."""
    RAW_DATA_DIR.mkdir(parents=True, exist_ok=True)
    EVALUATION_DIR.mkdir(parents=True, exist_ok=True)


def create_employee_handbook_pdf() -> None:
    """Generate the employee handbook PDF."""

    output_path = RAW_DATA_DIR / "employee_handbook.pdf"

    pdf = SimpleDocTemplate(
        str(output_path),
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()

    story = []

    story.append(
        Paragraph(
            "NexaTech Employee Handbook",
            styles["Title"],
        )
    )

    story.append(Spacer(1, 18))

    sections = [
        (
            "Annual Leave",
            (
                "Full-time employees are entitled to 25 days of paid "
                "annual leave per calendar year. Employees may carry "
                "forward a maximum of 5 unused annual leave days into "
                "the following calendar year."
            ),
        ),
        (
            "Probation Period",
            (
                "All new permanent employees complete a probation "
                "period of 6 months. A formal review is held before "
                "the end of the probation period."
            ),
        ),
        (
            "Sick Leave",
            (
                "Employees must inform their line manager before "
                "9:30 AM on the first day of sickness absence. "
                "Medical certification is required for absences "
                "lasting more than 7 calendar days."
            ),
        ),
        (
            "Learning and Development",
            (
                "Permanent employees receive an annual professional "
                "development allowance of £750. The allowance may be "
                "used for approved courses, certifications, books, "
                "and professional conferences."
            ),
        ),
    ]

    for heading, body in sections:
        story.append(
            Paragraph(
                heading,
                styles["Heading2"],
            )
        )

        story.append(
            Paragraph(
                body,
                styles["BodyText"],
            )
        )

        story.append(Spacer(1, 14))

    pdf.build(story)

    print(f"Created: {output_path.name}")


def create_remote_work_policy_docx() -> None:
    """Generate the remote working policy DOCX file."""

    output_path = RAW_DATA_DIR / "remote_work_policy.docx"

    document = WordDocument()

    document.add_heading(
        "NexaTech Remote Working Policy",
        level=0,
    )

    sections = [
        (
            "Remote Working Allowance",
            (
                "Eligible employees may work remotely for up to "
                "3 days per week. Remote working arrangements require "
                "approval from the employee's line manager."
            ),
        ),
        (
            "Eligibility",
            (
                "Employees become eligible for regular remote working "
                "after successfully completing their probation period."
            ),
        ),
        (
            "Company Equipment",
            (
                "NexaTech provides a company laptop, keyboard, mouse, "
                "and headset to permanent employees. Employees may "
                "request one external monitor for home working."
            ),
        ),
        (
            "Information Security",
            (
                "Employees must use the company VPN when accessing "
                "internal systems from outside a NexaTech office. "
                "Confidential company information must not be stored "
                "on personal devices."
            ),
        ),
        (
            "Working Abroad",
            (
                "Working remotely from another country requires "
                "written approval from both Human Resources and the "
                "employee's department director."
            ),
        ),
    ]

    for heading, body in sections:
        document.add_heading(
            heading,
            level=1,
        )
        document.add_paragraph(body)

    document.save(output_path)

    print(f"Created: {output_path.name}")


def create_company_faq_txt() -> None:
    """Generate the company FAQ TXT file."""

    output_path = RAW_DATA_DIR / "company_faq.txt"

    content = """NexaTech Company FAQ

OFFICE HOURS

Standard office hours are 9:00 AM to 5:30 PM,
Monday to Friday.

IT SUPPORT

The internal IT help desk is available from
8:00 AM to 6:00 PM on weekdays.

EXPENSE CLAIMS

Business expense claims must be submitted within
30 days of the date the expense was incurred.

BUILDING ACCESS

Employee access cards allow entry to NexaTech offices
between 7:00 AM and 10:00 PM.

PARKING

The Manchester office has 40 employee parking spaces.
Parking spaces must be reserved through the internal
booking system before arrival.

EMPLOYEE REFERRAL PROGRAMME

Employees receive a £500 referral bonus when a referred
candidate successfully completes their probation period.
"""

    output_path.write_text(
        content,
        encoding="utf-8",
    )

    print(f"Created: {output_path.name}")


def create_product_catalog_xlsx() -> None:
    """Generate the NexaTech product catalogue XLSX file."""

    output_path = RAW_DATA_DIR / "product_catalog.xlsx"

    workbook = Workbook()

    products_sheet = workbook.active
    products_sheet.title = "Products"

    products_sheet.append(
        [
            "Product",
            "Category",
            "Price_GBP",
            "Warranty_Years",
            "Stock",
        ]
    )

    products = [
        [
            "NexaBook Air",
            "Laptop",
            899,
            2,
            45,
        ],
        [
            "NexaBook Pro",
            "Laptop",
            1299,
            3,
            22,
        ],
        [
            "NexaStation Mini",
            "Desktop",
            749,
            2,
            18,
        ],
        [
            "NexaView 27",
            "Monitor",
            349,
            3,
            60,
        ],
        [
            "NexaView Ultra",
            "Monitor",
            599,
            4,
            15,
        ],
        [
            "NexaDock",
            "Accessory",
            129,
            1,
            80,
        ],
    ]

    for product in products:
        products_sheet.append(product)

    offices_sheet = workbook.create_sheet(
        title="Service Centres"
    )

    offices_sheet.append(
        [
            "City",
            "Support_Level",
            "Opening_Hours",
        ]
    )

    service_centres = [
        [
            "Manchester",
            "Full repair service",
            "08:30-17:30",
        ],
        [
            "London",
            "Full repair service",
            "09:00-18:00",
        ],
        [
            "Birmingham",
            "Drop-off only",
            "09:00-17:00",
        ],
    ]

    for centre in service_centres:
        offices_sheet.append(centre)

    workbook.save(output_path)

    print(f"Created: {output_path.name}")


def create_evaluation_dataset() -> None:
    """Create ground-truth questions for retrieval evaluation."""

    output_path = (
        EVALUATION_DIR
        / "evaluation_questions.json"
    )

    evaluation_questions = [
        {
            "question": (
                "How many days of annual leave do "
                "full-time employees receive?"
            ),
            "expected_answer": "25 days",
            "expected_source": "employee_handbook.pdf",
        },
        {
            "question": (
                "How many unused annual leave days can "
                "employees carry forward?"
            ),
            "expected_answer": "5 days",
            "expected_source": "employee_handbook.pdf",
        },
        {
            "question": (
                "How long is the probation period for "
                "new permanent employees?"
            ),
            "expected_answer": "6 months",
            "expected_source": "employee_handbook.pdf",
        },
        {
            "question": (
                "What is the annual professional development allowance?"
            ),
            "expected_answer": "£750",
            "expected_source": "employee_handbook.pdf",
        },
        {
            "question": (
                "How many days per week can eligible employees "
                "work remotely?"
            ),
            "expected_answer": "3 days per week",
            "expected_source": "remote_work_policy.docx",
        },
        {
            "question": (
                "What must employees use when accessing internal "
                "systems outside the office?"
            ),
            "expected_answer": "Company VPN",
            "expected_source": "remote_work_policy.docx",
        },
        {
            "question": (
                "When do employees become eligible for regular "
                "remote working?"
            ),
            "expected_answer": (
                "After successfully completing probation"
            ),
            "expected_source": "remote_work_policy.docx",
        },
        {
            "question": (
                "Within how many days must business expenses "
                "be submitted?"
            ),
            "expected_answer": "30 days",
            "expected_source": "company_faq.txt",
        },
        {
            "question": (
                "How many employee parking spaces does the "
                "Manchester office have?"
            ),
            "expected_answer": "40 parking spaces",
            "expected_source": "company_faq.txt",
        },
        {
            "question": (
                "What referral bonus is paid after a referred "
                "candidate completes probation?"
            ),
            "expected_answer": "£500",
            "expected_source": "company_faq.txt",
        },
        {
            "question": (
                "Which laptop has a three-year warranty?"
            ),
            "expected_answer": "NexaBook Pro",
            "expected_source": "product_catalog.xlsx",
        },
        {
            "question": (
                "How much does the NexaBook Pro cost?"
            ),
            "expected_answer": "£1299",
            "expected_source": "product_catalog.xlsx",
        },
        {
            "question": (
                "Which product has a four-year warranty?"
            ),
            "expected_answer": "NexaView Ultra",
            "expected_source": "product_catalog.xlsx",
        },
        {
            "question": (
                "Which service centre offers drop-off-only support?"
            ),
            "expected_answer": "Birmingham",
            "expected_source": "product_catalog.xlsx",
        },
    ]

    with output_path.open(
        "w",
        encoding="utf-8",
    ) as file:
        json.dump(
            evaluation_questions,
            file,
            indent=4,
            ensure_ascii=False,
        )

    print(f"Created: {output_path.name}")


def main() -> None:
    """Generate all synthetic project data."""

    ensure_directories()

    create_employee_handbook_pdf()
    create_remote_work_policy_docx()
    create_company_faq_txt()
    create_product_catalog_xlsx()
    create_evaluation_dataset()

    print("\nDataset generation complete.")


if __name__ == "__main__":
    main()