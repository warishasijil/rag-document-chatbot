from pathlib import Path

import pytest
from docx import Document as WordDocument
from openpyxl import Workbook
from reportlab.platypus import Paragraph, SimpleDocTemplate
from reportlab.lib.styles import getSampleStyleSheet


@pytest.fixture
def sample_txt(tmp_path: Path) -> Path:
    """Create a temporary TXT document."""

    file_path = tmp_path / "sample.txt"

    file_path.write_text(
        "Employees receive 25 days of annual leave.",
        encoding="utf-8",
    )

    return file_path


@pytest.fixture
def sample_docx(tmp_path: Path) -> Path:
    """Create a temporary DOCX document."""

    file_path = tmp_path / "sample.docx"

    document = WordDocument()

    document.add_heading(
        "Remote Working Policy",
        level=1,
    )

    document.add_paragraph(
        "Employees may work remotely three days per week."
    )

    document.save(file_path)

    return file_path


@pytest.fixture
def sample_xlsx(tmp_path: Path) -> Path:
    """Create a temporary Excel workbook."""

    file_path = tmp_path / "sample.xlsx"

    workbook = Workbook()

    worksheet = workbook.active
    worksheet.title = "Products"

    worksheet.append(
        [
            "Product",
            "Price_GBP",
            "Warranty_Years",
        ]
    )

    worksheet.append(
        [
            "NexaBook Pro",
            1299,
            3,
        ]
    )

    workbook.save(file_path)

    return file_path


@pytest.fixture
def sample_pdf(tmp_path: Path) -> Path:
    """Create a temporary single-page PDF."""

    file_path = tmp_path / "sample.pdf"

    styles = getSampleStyleSheet()

    pdf = SimpleDocTemplate(
        str(file_path)
    )

    pdf.build(
        [
            Paragraph(
                "Full-time employees receive "
                "25 days of annual leave.",
                styles["BodyText"],
            )
        ]
    )

    return file_path